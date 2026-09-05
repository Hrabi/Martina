#!/usr/bin/env python3
"""Build the first working DOCX export from the canonical manuscript manifest."""

from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


WORKING_TITLE_CS = (
    "Kontinuita ošetřovatelské péče u geriatrických pacientů při překladu "
    "z akutních oddělení na LDN: audit úplnosti a návaznosti sesterské "
    "překladové a přijímací dokumentace"
)
WORKING_TITLE_EN = "WORKING TITLE PENDING APPROVAL"
DRAFT_MARKER = "PRACOVNÍ KONCEPT - TESTOVACÍ EXPORT 2026-09-04"
START_PAGE_NUMBER = 7


def set_run_font(run, size: float | None = None, bold: bool | None = None, italic: bool | None = None):
    run.font.name = "Times New Roman"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.color.rgb = RGBColor(0, 0, 0)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_style_font(style, size: float, bold: bool = False, italic: bool = False):
    style.font.name = "Times New Roman"
    r_fonts = style._element.get_or_add_rPr().rFonts
    r_fonts.set(qn("w:ascii"), "Times New Roman")
    r_fonts.set(qn("w:hAnsi"), "Times New Roman")
    r_fonts.set(qn("w:eastAsia"), "Times New Roman")
    for attr in ("asciiTheme", "hAnsiTheme", "eastAsiaTheme", "cstheme"):
        r_fonts.attrib.pop(qn(f"w:{attr}"), None)
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.italic = italic
    style.font.color.rgb = RGBColor(0, 0, 0)


def set_page_geometry(section):
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(4)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)


def configure_styles(doc: Document):
    normal = doc.styles["Normal"]
    set_style_font(normal, 12)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.widow_control = True

    title = doc.styles["Title"]
    set_style_font(title, 16)
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.line_spacing = 1.15
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(18)
    title.paragraph_format.keep_with_next = True
    title_ppr = title._element.get_or_add_pPr()
    for border in title_ppr.findall(qn("w:pBdr")):
        title_ppr.remove(border)

    heading_specs = {
        "Heading 1": (14, True, 0),
        "Heading 2": (14, True, 24),
        "Heading 3": (12, True, 18),
        "Heading 4": (12, False, 12),
    }
    for style_name, (size, bold, before) in heading_specs.items():
        style = doc.styles[style_name]
        set_style_font(style, size, bold=bold)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(12 if style_name == "Heading 1" else 6)
        style.paragraph_format.line_spacing = 1.0
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    if "Pracovní poznámka" not in [style.name for style in doc.styles]:
        note = doc.styles.add_style("Pracovní poznámka", WD_STYLE_TYPE.PARAGRAPH)
    else:
        note = doc.styles["Pracovní poznámka"]
    set_style_font(note, 11, italic=True)
    note.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    note.paragraph_format.line_spacing = 1.15
    note.paragraph_format.space_after = Pt(8)


def add_field(paragraph, instruction: str):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text_node = OxmlElement("w:t")
    text_node.text = "7"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text_node, end])
    set_run_font(run, 12)


def set_page_number_start(section, start: int):
    sect_pr = section._sectPr
    pg_num_type = sect_pr.find(qn("w:pgNumType"))
    if pg_num_type is None:
        pg_num_type = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num_type)
    pg_num_type.set(qn("w:start"), str(start))


def set_update_fields(doc: Document):
    settings = doc.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def clear_story(story):
    p = story.paragraphs[0]
    for child in list(p._p):
        p._p.remove(child)
    return p


def configure_header(section):
    p = clear_story(section.header)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(DRAFT_MARKER)
    set_run_font(run, 8)


def configure_front_footer(section):
    p = clear_story(section.footer)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def configure_body_footer(section):
    section.footer.is_linked_to_previous = False
    p = clear_story(section.footer)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    add_field(p, "PAGE")


def add_centered_paragraph(doc: Document, text: str, size: float, *, bold=False, italic=False, after=0, before=0, style=None):
    p = doc.add_paragraph(style=style)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(text)
    set_run_font(run, size, bold=bold, italic=italic)
    return p


def add_heading_paragraph(doc: Document, text: str, level: int = 1):
    p = doc.add_paragraph(style=f"Heading {level}")
    sizes = {1: 14, 2: 14, 3: 12, 4: 12}
    bold = level <= 3
    run = p.add_run(text)
    set_run_font(run, sizes[level], bold=bold)
    return p


def add_title_page(doc: Document, *, english: bool):
    if english:
        institutions = [
            "UNIVERSITY OF OSTRAVA",
            "FACULTY OF MEDICINE",
            "DEPARTMENT OF NURSING AND MIDWIFERY",
        ]
        title = WORKING_TITLE_EN
        work_type = "RIGOROUS THESIS"
        author_label = "Author"
        consultant_label = "Consultant"
    else:
        institutions = [
            "OSTRAVSKÁ UNIVERZITA",
            "LÉKAŘSKÁ FAKULTA",
            "ÚSTAV OŠETŘOVATELSTVÍ A PORODNÍ ASISTENCE",
        ]
        title = WORKING_TITLE_CS
        work_type = "RIGORÓZNÍ PRÁCE"
        author_label = "Autor práce"
        consultant_label = "Konzultant"

    for idx, line in enumerate(institutions):
        add_centered_paragraph(doc, line, 14, bold=False, after=0 if idx < 2 else 118)

    p_title = doc.add_paragraph(style="Title")
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_after = Pt(18)
    run = p_title.add_run(title)
    set_run_font(run, 16, bold=False)

    add_centered_paragraph(doc, work_type, 14, after=126)

    values = [
        (f"{author_label}:", "údaj není potvrzen" if not english else "not confirmed"),
        (f"{consultant_label}:", "údaj není potvrzen" if not english else "not confirmed"),
    ]
    for label, value in values:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.tab_stops.add_tab_stop(Cm(4.5), WD_TAB_ALIGNMENT.LEFT)
        run = p.add_run(label)
        set_run_font(run, 12)
        p.add_run("\t")
        run = p.add_run(value)
        set_run_font(run, 12)

    add_centered_paragraph(doc, "2026", 12, after=0, before=72)


def parse_manuscript_manifest(path: Path) -> list[Path]:
    entries = []
    for line in path.read_text(encoding="utf-8").splitlines():
        match = re.match(r"\s*-\s+(drafts/.+\.md)\s*$", line)
        if match:
            entries.append(Path(match.group(1)))
    if not entries:
        raise ValueError(f"Manifest {path} does not contain manuscript pages")
    return entries


def parse_markdown(path: Path) -> tuple[dict[str, str], str]:
    text = path.read_text(encoding="utf-8")
    meta: dict[str, str] = {}
    if text.startswith("---\n"):
        _, raw_meta, text = text.split("---\n", 2)
        for line in raw_meta.splitlines():
            if ":" in line:
                key, value = line.split(":", 1)
                meta[key.strip()] = value.strip().strip('"')
    text = re.sub(r"<!--.*?-->", "", text, flags=re.S)
    lines = text.strip().splitlines()
    if lines and lines[0].startswith("# "):
        lines = lines[1:]
    return meta, "\n".join(lines).strip()


def clean_inline_markdown(text: str) -> str:
    text = re.sub(r"\[([^\]]+)\]\([^\)]+\)", r"\1", text)
    text = text.replace("`", "")
    text = text.replace("**", "").replace("__", "")
    return text.strip()


def add_markdown_body(doc: Document, body: str):
    if not body:
        p = doc.add_paragraph("Tato část zatím neobsahuje text.", style="Pracovní poznámka")
        return

    blocks = re.split(r"\n\s*\n", body)
    for block in blocks:
        block = block.strip()
        if not block:
            continue
        lines = block.splitlines()
        if all(re.match(r"^\s*[-*]\s+", line) for line in lines):
            for line in lines:
                text = clean_inline_markdown(re.sub(r"^\s*[-*]\s+", "", line))
                p = doc.add_paragraph(style="List Bullet")
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
                p.paragraph_format.space_after = Pt(0)
                run = p.add_run(text)
                set_run_font(run, 12)
            continue
        if len(lines) == 1 and re.match(r"^#{2,4}\s+", lines[0]):
            hashes, text = re.match(r"^(#{2,4})\s+(.+)$", lines[0]).groups()
            add_heading_paragraph(doc, clean_inline_markdown(text), level=len(hashes) - 1)
            continue
        text = clean_inline_markdown(" ".join(line.strip() for line in lines))
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.first_line_indent = Cm(0)
        run = p.add_run(text)
        set_run_font(run, 12)


def add_abstract_page(doc: Document, heading: str, body: str, keyword_label: str, keyword_text: str):
    p = add_heading_paragraph(doc, heading.upper(), level=1)
    p.paragraph_format.space_after = Pt(18)
    doc.add_paragraph("Stav části: osnova", style="Pracovní poznámka")
    add_markdown_body(doc, body)
    p_keywords = doc.add_paragraph()
    p_keywords.paragraph_format.space_before = Pt(18)
    p_keywords.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    run = p_keywords.add_run(keyword_label)
    set_run_font(run, 12, bold=True)
    run = p_keywords.add_run(keyword_text)
    set_run_font(run, 12)


def add_declaration_page(doc: Document):
    add_heading_paragraph(doc, "PROHLÁŠENÍ", level=1)
    doc.add_paragraph("Pracovní vzor - není určeno k podpisu", style="Pracovní poznámka")
    declaration_1 = (
        "Prohlašuji, že předložená práce je mým původním autorským dílem, které jsem "
        "vypracoval/a samostatně. Veškerou literaturu a další zdroje, z nichž jsem při "
        "zpracování čerpal/a, v práci řádně cituji a uvádím v seznamu použité literatury."
    )
    declaration_2 = (
        "Čestně prohlašuji, že text odevzdané práce v tištěné podobě je totožný s textem "
        "práce v elektronické podobě vložené do databáze DIPL2."
    )
    for text in (declaration_1,):
        p = doc.add_paragraph(text)
        p.paragraph_format.space_after = Pt(18)
    add_signature_block(doc)
    p = doc.add_paragraph(declaration_2)
    p.paragraph_format.space_before = Pt(150)
    p.paragraph_format.space_after = Pt(18)
    add_signature_block(doc)


def add_signature_block(doc: Document):
    table = doc.add_table(rows=2, cols=2)
    table.autofit = False
    table.columns[0].width = Cm(7)
    table.columns[1].width = Cm(7)
    for row in table.rows:
        row.cells[0].width = Cm(7)
        row.cells[1].width = Cm(7)
    values = [("V Ostravě dne", "................................"), ("", "podpis")]
    for row, data in zip(table.rows, values):
        for idx, (cell, value) in enumerate(zip(row.cells, data)):
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if idx == 1 else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(value)
            set_run_font(run, 12)
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "nil")
        borders.append(element)
    tbl_pr.append(borders)


def add_toc_page(doc: Document, entries: list[tuple[str, int, bool]]):
    add_heading_paragraph(doc, "OBSAH", level=1)
    for title, page, bold in entries:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(5)
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.tab_stops.add_tab_stop(Cm(14), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
        run = p.add_run(f"{title}\t{page}")
        set_run_font(run, 12, bold=bold)


def body_title(path: Path, source_title: str) -> str:
    mapping = {
        "10-introduction": "ÚVOD",
        "20-theoretical-background": "1 TEORETICKÁ VÝCHODISKA",
        "30-aims-and-questions": "2 CÍLE A VÝZKUMNÉ OTÁZKY",
        "40-methodology": "3 METODIKA",
        "50-results": "4 VÝSLEDKY",
        "60-discussion": "5 DISKUSE",
        "70-conclusion": "ZÁVĚR",
        "80-references": "SEZNAM POUŽITÝCH ZDROJŮ",
        "90-appendices": "PŘÍLOHY",
    }
    for key, title in mapping.items():
        if key in path.as_posix():
            return title
    return source_title.upper()


def build_document(repo_root: Path, output_path: Path):
    manifest_path = repo_root / "drafts" / "_manuscript.yml"
    manifest_entries = parse_manuscript_manifest(manifest_path)
    resolved_entries = [(path, repo_root / path) for path in manifest_entries]
    missing = [str(path) for path, resolved in resolved_entries if not resolved.exists()]
    if missing:
        raise FileNotFoundError(f"Missing manuscript pages: {', '.join(missing)}")

    pages = []
    for rel_path, abs_path in resolved_entries:
        meta, body = parse_markdown(abs_path)
        pages.append({"path": rel_path, "meta": meta, "body": body})

    doc = Document()
    set_page_geometry(doc.sections[0])
    configure_styles(doc)
    configure_header(doc.sections[0])
    configure_front_footer(doc.sections[0])
    set_update_fields(doc)

    doc.core_properties.title = "Testovací export rigorózní práce"
    doc.core_properties.subject = "Pracovní koncept sestavený z drafts/_manuscript.yml"
    doc.core_properties.author = ""
    doc.core_properties.last_modified_by = ""
    doc.core_properties.comments = "Pracovní testovací export, nikoli finální rukopis."

    abstract_cs = next(page for page in pages if "10-abstract-cs.md" in page["path"].as_posix())
    abstract_en = next(page for page in pages if "20-abstract-en.md" in page["path"].as_posix())
    body_pages = [
        page for page in pages
        if page["path"].name == "index.md" and "00-front-matter" not in page["path"].as_posix()
    ]

    add_title_page(doc, english=False)
    doc.add_page_break()
    add_title_page(doc, english=True)
    doc.add_page_break()
    add_abstract_page(doc, "Abstrakt", abstract_cs["body"], "Klíčová slova: ", "dosud nestanovena")
    doc.add_page_break()
    add_abstract_page(doc, "Abstract", abstract_en["body"], "Keywords: ", "not yet determined")
    doc.add_page_break()
    add_declaration_page(doc)
    doc.add_page_break()

    toc_entries = []
    for offset, page in enumerate(body_pages):
        title = body_title(page["path"], page["meta"].get("title", page["path"].stem))
        toc_entries.append((title, START_PAGE_NUMBER + offset, title.startswith(tuple(str(i) for i in range(1, 10))) or title in {"ÚVOD", "ZÁVĚR"}))
    add_toc_page(doc, toc_entries)

    body_section = doc.add_section(WD_SECTION.NEW_PAGE)
    set_page_geometry(body_section)
    body_section.header.is_linked_to_previous = True
    set_page_number_start(body_section, START_PAGE_NUMBER)
    configure_body_footer(body_section)

    for idx, page in enumerate(body_pages):
        if idx > 0:
            doc.add_page_break()
        title = body_title(page["path"], page["meta"].get("title", page["path"].stem))
        add_heading_paragraph(doc, title, level=1)
        doc.add_paragraph("Stav části: osnova", style="Pracovní poznámka")
        add_markdown_body(doc, page["body"])

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--output", required=True, type=Path)
    args = parser.parse_args()
    repo_root = Path(__file__).resolve().parents[1]
    build_document(repo_root, args.output.resolve())
    print(args.output.resolve())


if __name__ == "__main__":
    main()
